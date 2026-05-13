"""Generate PRD_VinhKhanhTour.docx — a full Product Requirements Document.

Cấu trúc:
  Trang bìa, Lịch sử phiên bản, Mục lục
  PHẦN A — SẢN PHẨM (executive summary, problem, goals/KPIs, personas, scope, pricing)
  PHẦN B — TRẢI NGHIỆM (user journeys, use case tổng thể)
  PHẦN C — KỸ THUẬT (kiến trúc, ERD, state machine, timing)
  PHẦN D — CHỨC NĂNG (F01-F10 + sub-flow F03b)
  PHẦN E — PHI CHỨC NĂNG (NFR, bảo mật, i18n)
  PHẦN F — API REFERENCE
  PHẦN G — VẬN HÀNH (rủi ro, DoD, giả định, lịch trình)
  PHẦN H — PHỤ LỤC (thuật ngữ, tham chiếu)
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
DIAGRAM_PNG_DIR = BASE_DIR / "diagrams" / "png"
OUT_DOCX = BASE_DIR / "PRD_VinhKhanhTour.docx"

PRODUCT_VERSION = "1.5.3"
DOC_VERSION = "2.2"
DOC_DATE = datetime.now().strftime("%d/%m/%Y")


# ──────────────────────────────────────────────────────────────────────────────
# Dữ liệu nguồn
# ──────────────────────────────────────────────────────────────────────────────

VERSION_HISTORY = [
    ["1.0", "17/04/2026", "Bản PRD ban đầu, phủ F01–F10."],
    ["1.5", "01/05/2026", "Đồng bộ sơ đồ với code; bổ sung dashboard."],
    ["1.5.1", "07/05/2026", "Refactor 3 thẻ dashboard (Tổng POI, POI hoạt động, Ngôn ngữ)."],
    ["2.0", "08/05/2026", "Tái cấu trúc PRD đầy đủ chuẩn: tóm tắt điều hành, KPI, personas, NFR, API ref, rủi ro, glossary. Đồng bộ với code v1.5.2 (helper CalculateRemainingDays)."],
    ["2.1", "09/05/2026", "Bổ sung bảng phân lớp triển khai và method names cho F01-F05; đồng bộ DOCX với PRD Markdown."],
    [DOC_VERSION, DOC_DATE, "Đồng bộ baseline kiểm thử: tests folder, API/CMS WebApplicationFactory, Playwright CMS smoke, Appium opt-in smoke và script run-all-tests."],
]


EXECUTIVE_SUMMARY = (
    "VinhKhanhTour là hệ thống du lịch ẩm thực thông minh phục vụ phố ẩm thực Vĩnh Khánh (Quận 4, TP.HCM). "
    "Sản phẩm gồm ba thành phần: ứng dụng MAUI cho khách du lịch, REST API ASP.NET Core làm backend trung tâm, "
    "và CMS Razor Pages cho quản trị viên. Khách di chuyển trong khu phố sẽ được thiết bị tự động phát thuyết minh "
    "đa ngôn ngữ (Việt/Anh/Trung) khi vào vùng geofence của từng quán, đồng thời CMS theo dõi vị trí khách thực tế "
    "và quản lý phí duy trì hằng tháng cho chủ quán."
)

PROBLEM_STATEMENT = [
    "Khách du lịch — đặc biệt khách quốc tế — không có thông tin tổng quan về phố ẩm thực bằng tiếng Anh/Trung; "
    "tour guide truyền thống đắt tiền và không co dãn theo nhịp đi của khách.",
    "Chủ quán không có kênh quảng bá nội dung quán (giới thiệu, thực đơn, audio thuyết minh) tới đúng khách đang đi ngang.",
    "Đơn vị quản lý phố không có công cụ giám sát số lượng và phân bố khách thực tế theo thời gian, gây khó khăn "
    "khi điều phối an ninh, chiếu sáng, vệ sinh.",
]

GOALS_KPI = [
    ["Tự động hoá thuyết minh", "Khách bước vào geofence trong 30m → app phát audio < 5 giây.", "Audio TTFB ≤ 5s; cooldown 10 phút/POI."],
    ["Đa ngôn ngữ", "Hỗ trợ Việt/Anh/Trung cho mọi POI; auto-fallback `vi` khi thiếu bản dịch.", "≥ 95% POI có đủ 3 bản dịch."],
    ["Theo dõi khách realtime", "CMS biết được thiết bị nào đang online trong vòng 2 phút.", "≥ 90% heartbeat thành công; độ trễ map ≤ 30s."],
    ["Mô hình kinh doanh hai chiều", "Khách trả phí gói; chủ quán trả phí duy trì hàng tháng.", "Tỷ lệ duyệt thanh toán < 24h: ≥ 95%."],
    ["Vận hành ổn định", "Triển khai 24/7, có cơ chế retry/fallback rõ ràng.", "Uptime API ≥ 99.0%; CMS dashboard load ≤ 3s."],
]


PERSONAS = [
    {
        "name": "Khách du lịch quốc tế",
        "role": "App User (anonymous)",
        "demographics": "20–40 tuổi, du lịch theo nhóm 2–4 người, không nói tiếng Việt.",
        "goals": [
            "Hiểu được nguồn gốc và đặc trưng từng món ăn của các quán dọc phố.",
            "Có hướng dẫn rảnh tay khi đi bộ — không phải đọc nhiều.",
            "Biết được thực đơn và mức giá trước khi vào quán.",
        ],
        "pain": [
            "Không có biển hiệu tiếng Anh/Trung tại phố.",
            "Không tìm được thông tin tập trung trên Google Maps.",
            "Hỏi nhân viên quán bằng tiếng Anh thường không thông.",
        ],
        "tech": "iPhone/Android tầm trung, dùng 4G; quen với QR và app du lịch.",
    },
    {
        "name": "Khách du lịch trong nước",
        "role": "App User (anonymous)",
        "demographics": "18–45 tuổi, đi nhóm bạn hoặc gia đình, có quen với app công nghệ.",
        "goals": [
            "Khám phá quán mới, biết quán nào đang đông khách thật.",
            "Đọc nhanh thực đơn để chọn quán phù hợp ngân sách.",
        ],
        "pain": [
            "Khó so sánh các quán liền kề khi không có review chung.",
            "Sợ bị 'chặt chém' khi chưa biết giá.",
        ],
        "tech": "Android phổ thông; quen QR thanh toán Vietnam (VietQR/Momo).",
    },
    {
        "name": "Quản trị viên hệ thống",
        "role": "Admin CMS",
        "demographics": "Nhân viên ban quản lý phố ẩm thực; làm việc giờ hành chính.",
        "goals": [
            "Quản lý nội dung POI, ảnh, thực đơn, thuyết minh đa ngôn ngữ.",
            "Duyệt yêu cầu thanh toán gói app trong vòng tối đa 24 giờ.",
            "Theo dõi bản đồ khách online và POI quá hạn duy trì.",
            "Ghi nhận phí duy trì hàng tháng và xuất hoá đơn.",
        ],
        "pain": [
            "Đối soát chuyển khoản thủ công với sao kê ngân hàng.",
            "Cần nhìn nhanh tổng quan để biết vấn đề phải xử lý.",
        ],
        "tech": "Trình duyệt Chrome/Edge trên Windows; quen Excel.",
    },
    {
        "name": "Chủ quán ăn",
        "role": "POI Owner",
        "demographics": "30–55 tuổi, làm việc trong quán, không rành công nghệ phức tạp.",
        "goals": [
            "Giới thiệu quán đến khách quốc tế qua thuyết minh đa ngôn ngữ.",
            "Cập nhật thực đơn, ảnh món, mức giá nhanh và đơn giản.",
            "Đóng phí duy trì để giữ POI hiển thị trên app.",
        ],
        "pain": [
            "Không có ngân sách tự xây website/app riêng.",
            "Mất khách quốc tế do rào cản ngôn ngữ.",
        ],
        "tech": "Android, đọc tin trên Zalo, dùng QR trả tiền điện nước.",
    },
]


IN_SCOPE = [
    "Mobile app MAUI Android: subscription gate, khám phá POI, geofence audio, chi tiết POI, thanh toán QR.",
    "REST API ASP.NET Core: POI, thuyết minh, heartbeat, subscription, payment, upload ảnh, log.",
    "CMS Razor Pages: dashboard, quản lý POI, duyệt thanh toán, ghi nhận phí duy trì, bản đồ live.",
    "Đa ngôn ngữ vi/en/zh cho POI, thuyết minh, UI app và CMS.",
    "Mô hình kinh doanh hai chiều (gói app cho khách, phí duy trì cho chủ quán).",
]

OUT_OF_SCOPE = [
    "Đặt bàn / order online từ app.",
    "Nền tảng iOS (App Store) — phiên bản 1.5.x chỉ phát hành Android.",
    "Push notification (Firebase / APNS).",
    "Đăng nhập mạng xã hội cho khách (sản phẩm đang dùng device-anonymous).",
    "Nhận diện hình ảnh món ăn / OCR menu giấy.",
    "Live chat khách – chủ quán.",
    "Đa thành phố hoặc đa khu vực phố ẩm thực — phiên bản hiện tại chỉ phục vụ một khu vực Vĩnh Khánh.",
]


PRICING_APP = [
    ["thu", "Dùng thử", "0đ", "3 ngày", "1 lần/thiết bị, kích hoạt ngay."],
    ["ngay", "1 ngày", "29.000đ", "1 ngày", "Thanh toán qua QR VietQR."],
    ["tuan", "1 tuần", "99.000đ", "7 ngày", "Thanh toán qua QR VietQR."],
    ["thang", "1 tháng", "199.000đ", "30 ngày", "Thanh toán qua QR VietQR."],
    ["nam", "1 năm", "999.000đ", "365 ngày", "Thanh toán qua QR VietQR."],
]

PRICING_POI = [
    ["Phí duy trì hằng tháng", "50.000đ/tháng", "Hằng tháng", "POI quá hạn → ẩn khỏi app cho đến khi gia hạn."],
    ["Phí convert TTS", "20.000đ/lần", "Mỗi lần convert", "Dùng khi chuyển nội dung text → audio."],
]


USER_JOURNEY_GUEST = [
    "1. Tải app, mở lần đầu → LaunchPage kiểm tra subscription cục bộ.",
    "2. Nếu chưa có gói → SubscriptionPage hiện 5 gói (thu/ngay/tuan/thang/nam) + recovery code.",
    "3. Chọn gói thu (miễn phí 3 ngày) → app kích hoạt ngay → vào MainPage.",
    "4. (Hoặc) Chọn gói trả phí → PaymentPage hiển thị QR VietQR + nội dung CK chuẩn.",
    "5. Khách chuyển khoản → bấm 'Đã chuyển khoản' → PaymentStatusPage poll trạng thái 10s/lần.",
    "6. Admin duyệt trên CMS → app nhận da_duyet → lưu hạn → vào MainPage.",
    "7. MainPage tải POI từ /api/poi, render danh sách + bản đồ Leaflet.",
    "8. Khách bật GPS, dạo phố. App lấy vị trí mỗi 5s.",
    "9. Khi vào geofence (BanKinh ≤ 30m), dwell 5s → confirm POI → enqueue audio.",
    "10. App đọc thuyết minh qua TTS theo ngôn ngữ đang chọn → ghi log vào /api/log.",
    "11. Khách bấm vào card POI → DetailPage hiện thực đơn, ảnh, nút Nghe lại, nút Chỉ đường.",
    "12. Khi gói sắp hết → app cảnh báo trong tab Cài đặt → khách quay lại bước 4 để gia hạn.",
]

USER_JOURNEY_ADMIN = [
    "1. Mở CMS → Dashboard hiện 3 thẻ (Tổng POI / POI đang hoạt động / Ngôn ngữ) + cảnh báo POI quá hạn.",
    "2. Vào /Poi để CRUD quán: tạo mới (tự gán hạn duy trì +30 ngày), upload ảnh, thêm thực đơn.",
    "3. Vào /ThuyetMinh quản lý nội dung audio guide cho từng POI, đa ngôn ngữ.",
    "4. Vào /BanDo theo dõi khách online, vị trí, POI hiện tại, XP và lịch sử hành trình 4h.",
    "5. Vào /DuyetThanhToan tab cho_duyet → kiểm tra nội dung CK với sao kê → bấm Duyệt/Từ chối.",
    "6. Vào /ThanhToan ghi nhận phí duy trì hằng tháng cho POI; mỗi tháng tạo 1 hóa đơn.",
    "7. Theo dõi tab thông báo POI quá hạn để liên hệ chủ quán.",
]


ARCHITECTURE_NOTES = [
    "Mobile app (MAUI Android) giao tiếp với REST API qua HTTPS; lưu device id trong Preferences.",
    "REST API (ASP.NET Core 10) là điểm trung gian duy nhất giữa app, CMS và database.",
    "CMS Razor Pages truy cập DB bằng EF Core (đa số) và raw Npgsql (analytics dashboard).",
    "DB Supabase PostgreSQL (managed cloud), kết nối qua connection string đặt trong appsettings/secrets.",
    "API publish trên Render (https://phoamthuc.onrender.com); CMS chạy nội bộ.",
    "Reverse map render bằng Leaflet (front-end) trên cả app (WebView) và CMS.",
]


DATA_TABLES = [
    ["poi", "Điểm du lịch / quán ăn", "KinhDo, ViDo, BanKinh (geofence m), MucUuTien, NgayHetHanDuyTri"],
    ["thuyetminh", "Nhóm nội dung audio guide của POI", "POIId (FK), ThuTu, TrangThai"],
    ["bandich", "Bản dịch nội dung thuyết minh", "ThuyetMinhId (FK), NgonNgu (vi/en/zh), NoiDung, FileAudio"],
    ["monan", "Món trong thực đơn quán", "POIId (FK), TenMonAn, DonGia, PhanLoai, HinhAnh, TinhTrang"],
    ["dangkydichvu", "Hợp đồng phí duy trì POI", "POIId (FK), PhiDuyTriThang, NgayBatDau, NgayHetHan, TrangThai"],
    ["hoadon", "Hóa đơn POI", "POIId, TaiKhoanId, LoaiPhi (duytri/convert), KyThanhToan (yyyy-MM)"],
    ["dangkyapp", "Gói đăng ký app của khách", "MaThietBi (logical), LoaiGoi, NgayBatDau, NgayHetHan, SoTien"],
    ["yeucauthanhtoan", "Yêu cầu thanh toán chờ duyệt", "MaThietBi, LoaiGoi, SoTien, NoiDungChuyen, TrangThai"],
    ["vitrikhach", "Vị trí GPS realtime của khách", "MaThietBi (UNIQUE), Lat, Lng, LanCuoiHeartbeat, PoiIdHienTai"],
    ["lichsuphat", "Audit log VIEW/GPS/app-geofence", "POIId, MaThietBi, Nguon, NgonNguDung, ThoiGian"],
    ["taikhoan", "Tài khoản admin/quản lý CMS", "TenDangNhap, MatKhau (BCrypt), VaiTro (admin/quanly)"],
]


TIMING_CONSTANTS = [
    ["GPS poll", "5 giây", "Geolocation.Default.GetLocationAsync() trong vòng lặp StartGpsTracking()."],
    ["Heartbeat", "10 giây", "POST /api/heartbeat (HEARTBEAT_EVERY_TICKS = 2)."],
    ["Refresh POI nền", "20 giây", "RefreshPoisInBackgroundAsync() (POI_REFRESH_EVERY_TICKS = 4)."],
    ["Dwell confirm POI", "5 giây", "Phải ở trong bán kính ≥ 5s mới enqueue audio + ghi visit."],
    ["Cooldown audio", "10 phút/POI", "_lastSpokenTime[poiKey] guard."],
    ["Visit dedup (server)", "10 phút", "POST /api/heartbeat/visit dedup ở DB."],
    ["View dedup (server)", "5 phút", "POST /api/heartbeat/view dedup ở DB."],
    ["Hysteresis 2 POI", "5 mét", "PoiSwitchDistanceBufferMeters — giữ _currentPoi nếu cùng MucUuTien và chênh ≤ 5m."],
    ["Online threshold", "2 phút", "VitriKhach.LanCuoiHeartbeat ≥ now - 2 phút → coi là 'online'."],
    ["Polling thanh toán", "10 giây", "PaymentStatusPage.PollStatusAsync()."],
    ["Snapshot duyệt CMS", "5 giây", "DuyetThanhToan ?handler=PendingSnapshot."],
]


NFR_PERFORMANCE = [
    ["API p95 latency", "≤ 500 ms", "PoiController.GetAll, HeartbeatController.SendHeartbeat."],
    ["Audio TTFB", "≤ 5 giây từ khi confirm POI", "Đo từ dwell-confirm → SpeakAsync."],
    ["CMS dashboard initial load", "≤ 3 giây", "Index.cshtml.cs LoadPoiSection + Analytics retry."],
    ["Geofence latency", "≤ 1 chu kỳ GPS (5s)", "Sau khi vào bán kính, app phải nhận diện trong vòng 1 lần poll."],
]

NFR_RELIABILITY = [
    "Heartbeat retry sạch: nếu API timeout, app giữ vị trí cũ và thử lại lần sau, không crash.",
    "Audio queue per-device có 3 lớp chống trùng: _playingPoi, _queuedSpeakPoiIds, _speakLock semaphore.",
    "Dashboard CMS có cơ chế retry ClearNpgsqlPoolsQuietly khi disposed wait handle.",
    "Subscription rollover: gia hạn nối tiếp NgayBatDau = NgayHetHan cũ → không mất ngày còn lại.",
    "Image fallback: ResolveImageUrl xử lý URL tương đối, localhost, URL ngoài, file null → ảnh placeholder.",
]

NFR_SECURITY = [
    "Không thu thập thông tin cá nhân khách: chỉ Device UUID tự sinh trong Preferences.",
    "CMS hiện chạy nội bộ; admin auth dùng BCrypt hash mật khẩu trong bảng `taikhoan`.",
    "API hiện chưa có JWT — phiên bản 1.5 chỉ filter qua MaThietBi (logical key) ở body request.",
    "Connection string Supabase và HostedApiBaseUrl đặt qua biến môi trường, không commit hard-code.",
    "Heartbeat verify tọa độ ngoài vùng phục vụ → server bỏ qua để không giả mạo 'đang ở quán'.",
    "VerifyCurrentPoiAsync: server tự re-check distance + 5m tolerance để chống spoofing claim POI.",
]

I18N_NOTES = [
    "Hệ ngôn ngữ: vi (mặc định), en, zh.",
    "App: MainPage chuyển ngôn ngữ qua nút trong tab Cài đặt; ThuyetMinh tự fallback về 'vi' khi thiếu bản dịch.",
    "CMS: dùng tiếng Việt; chưa hỗ trợ chuyển ngôn ngữ giao diện CMS.",
    "Mở rộng ngôn ngữ mới: chỉ cần chèn thêm BanDich.NgonNgu, không phải đổi schema.",
]

TEST_AUTOMATION = [
    ["API smoke/integration", "xUnit + Microsoft.AspNetCore.Mvc.Testing", "ApiHealthTests kiểm tra GET /health trả OK và có nội dung ok."],
    ["CMS smoke/E2E", "xUnit + Playwright Chromium", "CmsSmokeTests kiểm tra /health, mở /Privacy trong Chromium headless và đọc nội dung trang."],
    ["MAUI Appium smoke", "NUnit + Appium", "AppiumSmokeTests kiểm tra Appium /status khi RUN_APPIUM_TESTS=1; mặc định skip để không phụ thuộc emulator."],
    ["Script tổng hợp", "PowerShell/CMD", "scripts/run-all-tests.ps1 restore, build, cài Playwright Chromium, chạy API/CMS/Appium smoke; run-all-tests.cmd là wrapper Windows."],
]

TESTABILITY_NOTES = [
    "API và CMS khai báo `public partial class Program` để test project tạo host nội bộ bằng `WebApplicationFactory<Program>`.",
    "Test host dùng `ASPNETCORE_ENVIRONMENT=Testing` và connection string local `vinhkhanhtour_test`, không dùng production database.",
    "CMS E2E có thể tự start CMS tại `CMS_BASE_URL` (mặc định http://127.0.0.1:5199) nếu health check chưa sẵn sàng.",
    "Appium test là opt-in: chạy thật khi truyền `-WithAppium` cho script hoặc set `RUN_APPIUM_TESTS=1` sau khi bật emulator/device và Appium server.",
]


API_ENDPOINTS = {
    "POI": [
        ["GET", "/api/poi", "Danh sách POI đang hoạt động (TrangThai=true & NgayHetHanDuyTri ≥ now)."],
        ["GET", "/api/poi/{id}", "Chi tiết POI kèm thực đơn + thuyết minh theo ngôn ngữ."],
    ],
    "Subscription": [
        ["GET", "/api/subscription/plans", "Danh sách 5 gói (thu/ngay/tuan/thang/nam)."],
        ["GET", "/api/subscription/status/{maThietBi}", "Trạng thái gói hiện tại của thiết bị."],
        ["POST", "/api/subscription/purchase", "Mua gói thu miễn phí (kích hoạt ngay)."],
        ["POST", "/api/subscription/request", "Tạo yêu cầu thanh toán QR (cho_duyet)."],
        ["GET", "/api/subscription/request/{id}", "Polling trạng thái yêu cầu."],
        ["POST", "/api/subscription/approve/{id}", "Admin duyệt → tạo DangKyApp."],
        ["POST", "/api/subscription/reject/{id}", "Admin từ chối + lý do."],
    ],
    "Heartbeat & Tracking": [
        ["POST", "/api/heartbeat", "Upsert vị trí + POI hiện tại; verify lat/lng + 5m."],
        ["POST", "/api/heartbeat/visit", "Ghi visit khi vào geofence (dedup 10 phút)."],
        ["POST", "/api/heartbeat/view", "Ghi log mở chi tiết POI (dedup 5 phút)."],
        ["POST", "/api/heartbeat/sync-history", "Đồng bộ viewed/visited local nếu chênh."],
        ["GET", "/api/heartbeat/profile/{maThietBi}", "Hồ sơ XP, level, viewed/visited POI."],
        ["GET", "/api/heartbeat/active", "Thiết bị heartbeat trong 2 phút (CMS BanDo)."],
        ["GET", "/api/heartbeat/history/{deviceShort}", "Lịch sử POI 4 giờ gần nhất."],
    ],
    "Phí duy trì POI": [
        ["GET", "/api/payment/status/{poiId}", "Trạng thái hạn duy trì."],
        ["GET", "/api/payment/history/{poiId}", "Lịch sử hóa đơn POI."],
        ["POST", "/api/payment/maintenance", "Ghi nhận phí duy trì + tạo HoaDon."],
        ["GET", "/api/payment/overdue", "POI quá hạn duy trì."],
    ],
    "Khác": [
        ["GET", "/api/thuyet-minh/{poiId}?lang=vi", "Nội dung thuyết minh theo ngôn ngữ + fallback."],
        ["POST", "/api/log", "Ghi log phát audio (Nguon=GPS)."],
        ["POST", "/api/upload", "Upload ảnh ≤ 5MB (jpg/png/webp/gif)."],
        ["POST", "/api/auth/login", "Đăng nhập CMS (TenDangNhap + MatKhau BCrypt)."],
        ["POST", "/api/auth/register", "Đăng ký tài khoản admin/quản lý."],
        ["GET", "/health", "Health check nhanh cho API/CMS; dùng bởi smoke tests và probes."],
        ["GET", "/health/db", "Kiểm tra kết nối database cho API/CMS."],
    ],
}


RISKS = [
    ["RISK-001", "Cao", "Concurrency duyệt thanh toán: 2 admin Approve cùng 1 yêu cầu → tạo 2 DangKyApp.", "Thêm transaction + select for update; hoặc unique (YeuCauId) trong dangkyapp."],
    ["RISK-002", "Cao", "Polling thanh toán 404/500 → app im lặng vô hạn.", "Thêm backoff + max retry; UI hiển thị thông báo lỗi sau N lần fail."],
    ["RISK-003", "Trung bình", "Recovery code ghi đè trước khi server xác nhận có gói active.", "Server validate trước khi cho ghi override; rollback nếu sync fail."],
    ["RISK-004", "Trung bình", "GPS jitter giữa 2 POI cùng priority gây nháy POI.", "Đã giảm thiểu bằng hysteresis 5m + dwell 5s; cần test thực địa."],
    ["RISK-005", "Trung bình", "Render dashboard chậm khi DB nhiều dữ liệu (FillMissingBuckets, raw SQL)", "Tối ưu index trên lichsuphat(poiid, thoigian); thêm cache phía CMS."],
    ["RISK-006", "Thấp", "OneDrive lock file build artifact gây fail aapt2.", "Build trong path ngắn ngoài OneDrive; hoặc loại trừ thư mục obj/bin khỏi sync."],
    ["RISK-007", "Cao", "API public chưa có JWT — kẻ xấu có thể gọi endpoint nội bộ.", "Thêm middleware xác thực admin cho nhóm /api/subscription/approve, /api/payment/*."],
]

ACCEPTANCE_CRITERIA = [
    "Khách cài app, chọn gói thu → trong 3 giây có thể thấy MainPage với danh sách POI.",
    "Khách bước vào geofence một POI mới → audio thuyết minh phát trong ≤ 5 giây sau dwell 5 giây.",
    "Khách chuyển khoản và admin duyệt → app nhận trạng thái da_duyet trong ≤ 30 giây (tối đa 3 chu kỳ poll).",
    "Admin tạo POI mới → POI tự gán NgayHetHanDuyTri = now + 30 ngày, hiển thị ngay trên app sau lần refresh kế tiếp.",
    "POI quá hạn duy trì → bị ẩn khỏi GET /api/poi và CMS dashboard cảnh báo.",
    "Heartbeat liên tục thành công 95% trong điều kiện 4G ổn định; CMS hiện thiết bị 'online' trong 2 phút sau heartbeat cuối.",
    "Audio queue per-device không chồng chéo: TTS chỉ phát 1 POI tại 1 thời điểm; trùng POI bị skip.",
    "CMS Dashboard 3 thẻ: Tổng POI (đếm tất cả), POI đang hoạt động (TrangThai+còn hạn), Ngôn ngữ (3).",
    "Bộ smoke test tự động chạy được qua scripts/run-all-tests.ps1; Appium smoke được skip mặc định nếu chưa bật RUN_APPIUM_TESTS=1.",
]

ASSUMPTIONS = [
    "Khách có thiết bị Android 5.0+ với GPS bật và quyền vị trí được cấp.",
    "Mạng 4G/Wi-Fi hoạt động liên tục dọc phố Vĩnh Khánh; không có hầm/khu chết sóng.",
    "Hạ tầng Supabase PostgreSQL không bị downtime > 1 giờ/tháng.",
    "Render.com host API ở khu vực Singapore — ping từ Việt Nam ≤ 200ms.",
    "Người dùng đồng ý chia sẻ vị trí GPS để app hoạt động đúng.",
    "Admin có sao kê ngân hàng MBBank để đối soát yêu cầu thanh toán.",
]

CONSTRAINTS = [
    "Phiên bản 1.5.x chỉ phát hành Android (không iOS).",
    "Chỉ phục vụ một khu vực phố Vĩnh Khánh, Q4, TP.HCM (vùng phục vụ hard-code trong server).",
    "QR thanh toán dùng VietQR MBBank — chưa hỗ trợ ngân hàng khác.",
    "TTS dùng Microsoft.Maui.Media.TextToSpeech — phụ thuộc engine TTS của thiết bị.",
    "Đa ngôn ngữ chỉ có vi/en/zh; mở rộng thêm cần thay đổi config UI app/CMS.",
]

ROADMAP = [
    ["v1.0", "Q1/2026", "PRD ban đầu, dựng prototype POI list + audio TTS đơn giản."],
    ["v1.4.5", "03/2026", "Tích hợp GPS heartbeat, CMS BanDo, duyệt thanh toán."],
    ["v1.5.0", "04/2026", "Subscription gate hoàn chỉnh, recovery code/QR, dashboard CMS."],
    ["v1.5.1", "05/2026", "Refactor 3 thẻ dashboard; sub-flow F03b queue audio."],
    ["v1.5.2", "08/05/2026", "PRD bản đầy đủ chuẩn, đồng bộ feature F01-F10."],
    ["v1.5.3", DOC_DATE, "Bổ sung baseline kiểm thử tự động: API health, CMS Playwright smoke, Appium opt-in smoke, script run-all-tests."],
    ["v1.6 (kế hoạch)", "Q3/2026", "Thêm JWT auth cho API; iOS preview build; push notification."],
    ["v2.0 (định hướng)", "2027", "Đa khu vực phố ẩm thực; đặt bàn online; tích hợp Momo/ZaloPay."],
]

GLOSSARY = [
    ["POI", "Point of Interest — quán ăn / điểm tham quan trên bản đồ."],
    ["Geofence", "Vùng tròn quanh POI có bán kính BanKinh (mét). Vào trong → kích hoạt audio."],
    ["Dwell time", "Thời gian phải ở liên tục trong geofence trước khi confirm POI (5 giây)."],
    ["Hysteresis", "Cơ chế giữ POI hiện tại khi có 2 POI sát nhau, để tránh nháy do GPS jitter (5m buffer)."],
    ["Cooldown", "Khoảng thời gian không phát lại audio cho cùng 1 POI (10 phút/POI)."],
    ["Heartbeat", "Tín hiệu định kỳ app gửi server để chứng tỏ thiết bị còn sống (mỗi 10 giây)."],
    ["MaThietBi", "Mã thiết bị logic, sinh từ Preferences app, dùng làm khóa định danh khách ẩn danh."],
    ["Recovery code", "Chuỗi mã hóa MaThietBi để khách khôi phục thiết bị khi đổi máy / cài lại app."],
    ["TTS", "Text-to-Speech: chuyển văn bản thành giọng đọc."],
    ["TTFB", "Time To First Byte — thời gian từ request đến byte đầu tiên của response."],
    ["XP / Level", "Điểm trải nghiệm tích lũy theo viewed*50 + visited*100; level mỗi 500 XP."],
    ["VietQR", "Tiêu chuẩn QR thanh toán liên ngân hàng Việt Nam."],
    ["TrangThai cho_duyet/da_duyet/tu_choi", "3 trạng thái của YeuCauThanhToan."],
]


# ──────────────────────────────────────────────────────────────────────────────
# Cấu hình use case + features (giữ nguyên data cũ)
# ──────────────────────────────────────────────────────────────────────────────

USE_CASE_SUMMARY = {
    "title": "Use case tổng thể hệ thống",
    "image": "00-overall-usecase.png",
    "description": (
        "Sơ đồ use case tổng thể mô tả toàn bộ phạm vi chức năng đang có trong sản phẩm. "
        "Các chức năng chính đã phủ hết các màn hình MAUI, các controller API đang được gọi thật, "
        "và các Razor PageModel trong CMS."
    ),
    "coverage_note": (
        "Sau khi đối chiếu lại code hiện tại, chưa phát hiện thêm chức năng nghiệp vụ độc lập nào "
        "cần tách thành một use case mới ngoài F01 đến F10. Các luồng như khôi phục gói bằng mã QR, "
        "khôi phục mã thiết bị, polling, upload ảnh, geofence và analytics đều đã được mô hình hóa "
        "như chức năng chính hoặc subflow của các feature hiện có."
    ),
}


FEATURES = [
    {
        "id": "F01",
        "title": "Khởi động, kích hoạt và khôi phục quyền sử dụng app",
        "actor": "Khách du lịch",
        "status": "VERIFIED",
        "activity_image": "01-subscription-gate-activity.png",
        "sequence_image": "01-subscription-gate-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả chức năng mở ứng dụng, kiểm tra gói đang có, chặn người dùng khi chưa đủ điều kiện sử dụng, "
            "và cho phép khôi phục gói cũ bằng recovery code hoặc mã QR."
        ),
        "capabilities": [
            "Điều hướng khởi động qua LaunchPage rồi quyết định vào MainPage hay SubscriptionPage.",
            "Kiểm tra subscription cục bộ trước, sau đó thử phục hồi trạng thái gói từ server nếu dữ liệu local không còn hợp lệ.",
            "Cho phép dùng gói dùng thử hoặc chuyển sang luồng thanh toán gói trả phí.",
            "Hiển thị recovery payload và QR để người dùng sao chép, nhập tay, dán từ clipboard hoặc quét mã QR để khôi phục thiết bị cũ.",
        ],
        "implementation": [
            ["UI (boundary)", "`LaunchPage.xaml`, `SubscriptionPage.xaml`, `QrScannerPage.xaml`."],
            ["Page/Control", "`LaunchPage.OnAppearing / RouteAsync`, `MainPage.OnAppearing / EnsureSubscriptionGateAsync`, `SubscriptionPage.UpdateRecoveryCard`, `OnMuaGoiClicked`, `OnPasteRecoveryCodeClicked`, `OnEnterRecoveryCodeClicked`, `OnScanRecoveryQrClicked`, `RestoreFromRecoveryCodeAsync`, `RestoreSubscriptionStateAsync`, `ActivateFreeTrialAsync`."],
            ["Domain/helper", "`SubscriptionState.IsSubscriptionActive`, `HasStoredSubscriptionRecord`, `CalculateRemainingDays`; `DeviceIdentity.GetDeviceId`, `BuildRecoveryPayload`, `BuildQrCodeUrl`, `TrySetDeviceIdOverride`; `ApiConnectionPrompt.EnsureConnectedApiBaseUrlAsync`."],
            ["Controller", "`SubscriptionController.GetStatus` (`GET /api/subscription/status/{maThietBi}`), `Purchase` (`POST /api/subscription/purchase`)."],
            ["Persistence", "`Preferences` (`sub_ngay_het_han`, `da_dung_thu`, `device_id`, `device_id_override`) + `AppDbContext` -> PostgreSQL bảng `dangkyapp`."],
        ],
        "operation": [
            "Khi app mở, LaunchPage chờ ngắn rồi gọi RouteAsync để kiểm tra SubscriptionState.",
            "Nếu local subscription còn hạn, app đi thẳng vào MainPage; nếu không thì mở SubscriptionPage.",
            "Trong SubscriptionPage, người dùng có thể sao chép recovery code hoặc quét QR; app yêu cầu quyền camera trước khi mở QrScannerPage.",
            "Khi nhận được mã hợp lệ, app gọi TrySetDeviceIdOverride(), sau đó GET /api/subscription/status/{deviceId} để khôi phục trạng thái gói.",
            "Server trả về CoDangKy, NgayHetHan, SoNgayConLai (tính bằng Math.Floor + min 0 từ v1.5.2 qua helper SubscriptionState.CalculateRemainingDays), LoaiGoi, DaDungThu.",
            "Nếu server trả về gói còn hạn, app cập nhật Preferences, cập nhật nút dùng thử và đóng subscription gate.",
        ],
    },
    {
        "id": "F02",
        "title": "Khám phá danh sách POI, đồng bộ lịch sử và khôi phục mã thiết bị",
        "actor": "Khách du lịch",
        "status": "VERIFIED",
        "activity_image": "02-poi-explore-activity.png",
        "sequence_image": "02-poi-explore-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả chức năng tải danh sách quán ăn/POI, tìm kiếm, mở trang chi tiết, đồng bộ lịch sử xem/thăm, "
            "và khôi phục lại mã thiết bị ngay trong tab cài đặt của MainPage."
        ),
        "capabilities": [
            "Tải danh sách POI từ API và dùng fallback data khi API lỗi.",
            "Đồng bộ hồ sơ trải nghiệm, XP, viewed/visited POI từ /api/heartbeat/profile và /api/heartbeat/sync-history.",
            "Hỗ trợ tìm kiếm trực tiếp và render lại card/map theo từ khóa.",
            "Cho phép sao chép hoặc khôi phục mã thiết bị bằng nhập tay, clipboard hoặc quét QR trong tab cài đặt.",
        ],
        "implementation": [
            ["UI (boundary)", "`MainPage.xaml` (tab Khám phá/Bản đồ/Cài đặt, search box, POI cards), `QrScannerPage.xaml`."],
            ["Page/Control", "`MainPage.OnAppearing`, `LoadPoisFromApi`, `RefreshPoisInBackgroundAsync`, `RenderPoiCards`, `OnSearchTextChanged`, `OpenPoiDetailAsync`, `RestorePoiHistoryAsync`, `SyncPoiHistoryAsync`, `UpdateCaiDatUI`, `RecordPoiViewAsync`."],
            ["Domain/helper", "`DeviceIdentity.GetDeviceId`, `GetSavedPoiIds`, `MergeSavedPoiIds`, `FoodImageCatalog.GetPoiImageSource`, `AppConfig.EnsureApiBaseUrlAsync`."],
            ["Controller", "`PoiController.GetAll` (`GET /api/poi`), `HeartbeatController.GetExperienceProfile`, `SyncHistory`, `RecordView`."],
            ["Persistence", "`Preferences` (viewed/visited POI ids, device id) + `AppDbContext` -> PostgreSQL bảng `poi`, `lichsuphat`."],
        ],
        "operation": [
            "Khi MainPage xuất hiện lần đầu, app chạy LoadPoisFromApi() rồi gọi GET /api/poi.",
            "Sau khi có dữ liệu, app phục hồi lịch sử từ server, đồng bộ lại local history và render card/map.",
            "Mỗi 20 giây app có thể refresh danh sách POI nền để dữ liệu hiển thị không quá cũ.",
            "Khi người dùng mở chi tiết POI, app ghi viewed local, POST /api/heartbeat/view, rồi đồng bộ lịch sử.",
            "Trong tab cài đặt, khi người dùng nhập hoặc quét mã QR cũ, app gọi RestoreDeviceCodeAsync() để khôi phục subscription, lịch sử xem/thăm và mã QR mới của thiết bị.",
            "Hồ sơ XP từ /api/heartbeat/profile bao gồm ConLaiNgay; từ v1.5.2 tính bằng Math.Floor + min 0 (đồng nhất với /api/subscription/status).",
        ],
    },
    {
        "id": "F03",
        "title": "Theo dõi GPS, geofence và tự động phát thuyết minh",
        "actor": "Khách du lịch",
        "status": "VERIFIED",
        "activity_image": "03-geofence-audio-activity.png",
        "sequence_image": "03-geofence-audio-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả phần lõi vận hành theo vị trí: lấy GPS, gửi heartbeat, xác định POI gần nhất, "
            "ghi nhận visit và tự động phát thuyết minh khi người dùng đi vào vùng geofence. "
            "Bao gồm xử lý đặc biệt khi khách đứng giữa 2 POI (sort theo MucUuTien → distance → tên + hysteresis 5m)."
        ),
        "capabilities": [
            "Lấy vị trí định kỳ và cập nhật trạng thái người dùng trên bản đồ.",
            "Gửi heartbeat vị trí về server để CMS biết thiết bị nào đang online.",
            "Xác định geofence theo bán kính từng POI và chống phát lặp bằng cooldown 10 phút.",
            "Hysteresis 5m + dwell 5s tránh nháy POI khi GPS jitter giữa 2 quán liền kề.",
            "Tự động đọc nội dung thuyết minh và ghi log phát audio.",
        ],
        "implementation": [
            ["UI (boundary)", "`MainPage.xaml` (GPS status, map WebView, now-playing banner, POI highlight)."],
            ["Page/Control", "`EnsureGpsTrackingAsync`, `StartGpsTracking`, `CheckGeofence`, `SelectPoiForLocation`, `QueuePoiPlayback`, `ProcessSpeakQueueAsync`, `SpeakPoiAsync`, `SendHeartbeatAsync`, `RecordPoiVisitAsync`, `LogPlaybackAsync`."],
            ["Domain/helper", "`Location.CalculateDistance`, `_lastSpokenTime`, `_speakQueue`, `_queuedSpeakPoiIds`, `_speakLock`, `AppText.LanguageCode`, `TextToSpeech.Default`."],
            ["Controller", "`HeartbeatController.SendHeartbeat`, `RecordVisit`, `ThuyetMinhController.GetByPoi`, `LogController.Post`."],
            ["Persistence", "`Preferences` (visited ids) + `AppDbContext` -> PostgreSQL bảng `vitrikhach`, `lichsuphat`, `thuyetminh`, `bandich`."],
        ],
        "operation": [
            "MainPage gọi EnsureGpsTrackingAsync(), xin quyền vị trí và bắt đầu vòng lặp GPS.",
            "App lấy vị trí mỗi 5 giây bằng Geolocation.Default.GetLocationAsync().",
            "SelectPoiForLocation() sort candidate theo MucUuTien → DistanceMeters → TenPOI; nếu POI hiện tại cùng priority và chỉ chênh ≤ 5m thì giữ.",
            "Heartbeat được gửi mỗi 10 giây dựa trên HEARTBEAT_EVERY_TICKS = 2; refresh POI nền mỗi 20 giây dựa trên POI_REFRESH_EVERY_TICKS = 4.",
            "Khi phát hiện người dùng vừa đi vào một POI mới và đã qua cooldown, app ghi visited local, POST /api/heartbeat/sync-history, POST /api/heartbeat/visit rồi cập nhật UI.",
            "Sau đó app lấy thuyết minh qua /api/thuyet-minh/{poiId}, đọc bằng TTS và POST /api/log để lưu lịch sử phát.",
        ],
    },
    {
        "id": "F04",
        "title": "Xem chi tiết POI, thực đơn, audio guide và chỉ đường",
        "actor": "Khách du lịch",
        "status": "VERIFIED",
        "activity_image": "04-poi-detail-activity.png",
        "sequence_image": "04-poi-detail-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả trang chi tiết quán ăn: tải nội dung POI, hiện ảnh bìa và món ăn, "
            "phát audio/thuyết minh và mở chỉ đường ngoài app."
        ),
        "capabilities": [
            "Tải chi tiết POI theo ngôn ngữ đang dùng.",
            "Dùng FoodImageCatalog để sinh ảnh fallback cho quán và món ăn khi URL ảnh lỗi hoặc thiếu.",
            "Phát audio file bằng AudioWebView khi có file, hoặc đọc text bằng TextToSpeech khi chỉ có nội dung chữ.",
            "Mở Google Maps / browser để chỉ đường đến quán.",
        ],
        "implementation": [
            ["UI (boundary)", "`DetailPage.xaml` (cover, info, menu, audio WebView, play/pause/stop/slider, nút chỉ đường)."],
            ["Page/Control", "`DetailPage.LoadDetail`, `ConfigureAudioPlayer`, `RenderMenu`, `UseFallback`, `OnNgheClicked`, `OnAudioWebViewNavigated`, `OnAudioWebViewNavigating`, `OnPlayPauseClicked`, `OnStopAudioClicked`, `OnAudioSliderDragCompleted`, `OnMapClicked`."],
            ["Domain/helper", "`FoodImageCatalog.GetPoiImageSource`, `TextToSpeech.Default`, `Browser.Default.OpenAsync`, `AppText.LanguageCode`."],
            ["Controller", "`PoiController.GetById` (`GET /api/poi/{id}?lang={lang}`), `HeartbeatController.RecordView` từ luồng mở chi tiết ở `MainPage.RecordPoiViewAsync`."],
            ["Persistence", "`AppDbContext` -> PostgreSQL bảng `poi`, `monan`, `thuyetminh`, `bandich`, `lichsuphat`."],
        ],
        "operation": [
            "DetailPage khởi tạo audio bridge, apply label rồi gọi LoadDetail().",
            "Nếu API GET /api/poi/{id}?lang=... thành công, trang render toàn bộ thông tin quán, menu và cấu hình player audio.",
            "Nếu API lỗi hoặc không trả về dữ liệu hợp lệ, trang chuyển sang UseFallback(tenPoi), dùng ảnh fallback và ẩn các vùng không đủ dữ liệu.",
            "Khi bấm Nghe, nếu có FileAudio thì AudioWebView thực hiện play/pause/stop/seek; nếu không có file thì dùng TextToSpeech đọc nội dung thuyết minh.",
            "Khi bấm Chỉ đường, app mở Google Maps bằng tọa độ của POI.",
        ],
    },
    {
        "id": "F05",
        "title": "Thanh toán gói trả phí và chờ duyệt",
        "actor": "Khách du lịch",
        "status": "VERIFIED",
        "activity_image": "05-paid-plan-activity.png",
        "sequence_image": "05-paid-plan-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả luồng thanh toán gói app bằng mã QR chuyển khoản, từ lúc tạo nội dung chuyển khoản "
            "đến lúc polling trạng thái duyệt và quay lại app."
        ),
        "capabilities": [
            "Sinh nội dung chuyển khoản dạng VKT <GOI> <SHORTID> và URL QR VietQR.",
            "Tạo yêu cầu thanh toán trên server qua /api/subscription/request.",
            "Poll trạng thái duyệt của yêu cầu để biết đã duyệt, từ chối hay còn chờ.",
            "Đóng modal stack đúng logic khi bắt đầu dùng hoặc quay lại thử thanh toán tiếp.",
        ],
        "implementation": [
            ["UI (boundary)", "`SubscriptionPage.xaml`, `PaymentPage.xaml`, `PaymentStatusPage.xaml`."],
            ["Page/Control", "`SubscriptionPage.OnMuaGoiClicked`, `PaymentPage.SetupUi`, `OnCopyNoiDungClicked`, `OnDaChuyenKhoanClicked`, `OnConfigureApiClicked`, `PaymentStatusPage.OnAppearing`, `StartPollingAsync`, `PollStatusAsync`, `ShowSuccess`, `ShowRejected`, `ClosePaymentFlowAsync`."],
            ["Domain/helper", "`DeviceIdentity.GetDeviceId`, `ApiConnectionPrompt.EnsureConnectedApiBaseUrlAsync`, `SubscriptionState.CalculateRemainingDays`, `AppConfig.EnsureApiBaseUrlAsync`."],
            ["Controller", "`SubscriptionController.CreateRequest` (`POST /api/subscription/request`), `GetRequestStatus` (`GET /api/subscription/request/{yeuCauId}`)."],
            ["Persistence", "`Preferences[\"sub_ngay_het_han\"]` + `AppDbContext` -> PostgreSQL bảng `yeucauthanhtoan`, `dangkyapp`."],
        ],
        "operation": [
            "Từ SubscriptionPage, người dùng chọn gói trả phí và mở PaymentPage.",
            "PaymentPage gọi DeviceIdentity.GetDeviceId(), dựng _noiDungChuyen và hiển thị QR chuyển khoản.",
            "Khi người dùng bấm Đã chuyển khoản, app POST /api/subscription/request để tạo yêu cầu chờ duyệt.",
            "Sau đó PaymentStatusPage bắt đầu polling GET /api/subscription/request/{id} mỗi 10 giây.",
            "Nếu duyệt thành công, app lưu ngày hết hạn vào Preferences và chạy ClosePaymentFlowAsync(closeSubscriptionPage:true); nếu bị từ chối thì hiển thị lý do và cho phép thử lại hoặc đóng.",
        ],
    },
    {
        "id": "F06",
        "title": "Quản lý POI, ảnh, geofence, thuyết minh và menu trong CMS",
        "actor": "Quản trị viên CMS",
        "status": "VERIFIED",
        "activity_image": "06-cms-poi-management-activity.png",
        "sequence_image": "06-cms-poi-management-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả chức năng quản trị nội dung quán ăn trong CMS, bao gồm danh sách POI, tạo/sửa quán, "
            "upload ảnh, cập nhật geofence, thuyết minh đa ngôn ngữ và thực đơn món ăn."
        ),
        "capabilities": [
            "Danh sách /Poi hỗ trợ search, status filter, expiry filter và sort.",
            "Upload ảnh cover hoặc ảnh món ăn qua /api/upload rồi lưu URL vào form.",
            "Tạo và sửa thuyết minh đa ngôn ngữ Việt/Anh/Trung.",
            "Quản lý danh sách MonAns với logic bỏ qua row rỗng, mặc định DonGia = 0 khi null, và bảo toàn input người dùng khi validation POI lỗi.",
        ],
        "operation": [
            "Admin có thể vào /Poi để xem toàn bộ quán, lọc trạng thái hiển thị và sắp xếp theo tiêu chí đang hỗ trợ.",
            "Ở trang Create/Edit, khi chọn file ảnh, JS gọi doUpload() để POST /api/upload rồi cập nhật preview ngay trên form.",
            "CreateModel.OnPostAsync() tạo mới POI, thuyết minh, bản dịch và danh sách món ăn hợp lệ từ form.",
            "EditModel.OnPostAsync() loại các lỗi validation trên MonAns trống, giữ lại phần user vừa nhập khi POI còn lỗi, đồng thời deactivate món không còn trong form và add/update món còn lại.",
            "ImageUrlHelper.ResolvePoi() và ResolveDish() được dùng trong view để hiển thị đúng ảnh thật hoặc ảnh fallback cho quán và món ăn.",
        ],
    },
    {
        "id": "F07",
        "title": "Ghi nhận phí duy trì và lịch sử hóa đơn POI",
        "actor": "Quản trị viên CMS",
        "status": "VERIFIED",
        "activity_image": "07-maintenance-payment-activity.png",
        "sequence_image": "07-maintenance-payment-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả nghiệp vụ thu phí duy trì hằng tháng cho quán ăn và tra cứu lại lịch sử hóa đơn của từng POI."
        ),
        "capabilities": [
            "Trang /ThanhToan tổng hợp trạng thái hạn duy trì, số quán đã đóng, số quán quá hạn, số quán sắp hết hạn và tổng thu tháng.",
            "Trang /ThanhToan/GhiNhan cho phép gia hạn theo số tháng và cập nhật đơn giá phí duy trì.",
            "Tạo một bản ghi HoaDon cho từng kỳ thanh toán khi gia hạn nhiều tháng.",
            "Trang /ThanhToan/LichSu cho phép xem đầy đủ lịch sử hóa đơn của từng POI.",
        ],
        "operation": [
            "IndexModel.OnGetAsync() tải danh sách POI, các gói DangKyDichVu active và doanh thu tháng hiện tại.",
            "Khi admin chọn một POI để ghi nhận phí, GhiNhanModel.LoadPoiAsync() nạp thông tin quán, gói active và 5 hóa đơn gần nhất.",
            "Khi submit, hệ thống tính mốc gia hạn mới dựa trên hạn cũ hoặc thời điểm hiện tại, cập nhật poi.NgayHetHanDuyTri và bảo đảm poi.TrangThai = true.",
            "Nếu chưa có DangKyDichVu active thì tạo mới; nếu đã có thì cập nhật PhiDuyTriThang và NgayHetHan.",
            "Sau đó hệ thống tạo một HoaDon cho mỗi tháng gia hạn rồi lưu toàn bộ thay đổi vào database.",
        ],
    },
    {
        "id": "F08",
        "title": "Duyệt hoặc từ chối thanh toán gói app",
        "actor": "Quản trị viên CMS",
        "status": "VERIFIED",
        "activity_image": "08-app-payment-approval-activity.png",
        "sequence_image": "08-app-payment-approval-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả luồng CMS dùng để xử lý các yêu cầu thanh toán gói app mà khách gửi từ mobile app."
        ),
        "capabilities": [
            "Tải danh sách yêu cầu theo tab chờ duyệt, đã duyệt và đã từ chối.",
            "Cập nhật badge và cảnh báo bằng AJAX polling qua handler PendingSnapshot.",
            "Duyệt yêu cầu để tạo DangKyApp mới và tính lại hạn sử dụng.",
            "Từ chối yêu cầu và lưu lý do cho người dùng xem lại ở phía app.",
        ],
        "operation": [
            "Khi trang /DuyetThanhToan mở, CMS tải thống kê toàn bộ yêu cầu và danh sách yêu cầu theo tab đang chọn.",
            "JavaScript trên trang gọi ?handler=PendingSnapshot mỗi 5 giây để cập nhật số lượng chờ duyệt và phát hiện yêu cầu mới.",
            "Nếu admin bấm Duyệt, OnPostApproveAsync() tìm gói tương ứng, tính mốc bắt đầu và hết hạn mới, tạo DangKyApp rồi cập nhật YeuCauThanhToan sang da_duyet.",
            "Nếu admin bấm Từ chối, OnPostRejectAsync() chuyển trạng thái sang tu_choi, lưu NgayDuyet và GhiChuAdmin.",
            "Kết quả xử lý được redirect về đúng tab để admin tiếp tục theo dõi các yêu cầu còn lại.",
        ],
    },
    {
        "id": "F09",
        "title": "Theo dõi khách hàng, trạng thái sử dụng, XP và hành trình",
        "actor": "Quản trị viên CMS",
        "status": "VERIFIED",
        "activity_image": "09-live-map-activity.png",
        "sequence_image": "09-live-map-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả trang BanDo trong CMS dùng để theo dõi thiết bị khách, trạng thái online, gói sử dụng, "
            "mức độ tương tác với POI và chỉ số trải nghiệm."
        ),
        "capabilities": [
            "Đọc dữ liệu thuê bao, vị trí hiện tại và lịch sử hoạt động bằng raw Npgsql query.",
            "Tính viewed POI, visited POI, XP, level, progress, total spent và số gói đã mua cho từng thiết bị.",
            "Hỗ trợ search, filter và sort server-side.",
            "Tạo badge trạng thái online, ở quán, hết hạn và helper mô tả thời gian còn lại.",
        ],
        "operation": [
            "BanDo/Index.OnGetAsync() chuẩn hóa filter/sort rồi nạp dữ liệu từ các truy vấn raw Npgsql có timeout và retry.",
            "Hệ thống đọc max hạn gói từ dangkyapp, vị trí từ vitrikhach, số POI đã ghé và đã xem từ lichsuphat.",
            "BuildCustomerRows() hợp nhất các nguồn dữ liệu theo deviceId để tạo ra trạng thái online, current POI, viewedCount, visitedCount và thông tin gói.",
            "XP được tính theo công thức viewed * 50 + visited * 100; level được suy ra với ExperiencePerLevel = 500.",
            "Trang BanDo hiện tại render server-side, chưa có chu kỳ auto-refresh cố định trong code; admin xem dữ liệu mới bằng lần tải trang kế tiếp.",
        ],
    },
    {
        "id": "F10",
        "title": "Dashboard tổng quan và monitoring CMS",
        "actor": "Quản trị viên CMS",
        "status": "VERIFIED",
        "activity_image": "10-dashboard-activity.png",
        "sequence_image": "10-dashboard-sequence.png",
        "function_text": (
            "Nhóm sơ đồ này mô tả dashboard tổng hợp của CMS, bao gồm danh sách POI, các chỉ số monitoring, heatmap hoạt động, "
            "top POI và doanh thu theo khoảng thời gian."
        ),
        "capabilities": [
            "Hỗ trợ nhiều chế độ khoảng thời gian: today, last7, last30, last12m, day, week, month, year và custom.",
            "Tách riêng phần tải danh sách POI và phần analytics có retry.",
            "Đọc activity, geo heat, top POI, revenue và summary bằng raw Npgsql trên shared EF connection.",
            "Bù bucket thiếu bằng FillMissingBuckets() để biểu đồ không bị thủng mốc thời gian.",
            "3 thẻ thống kê chính: Tổng POI (tất cả), POI đang hoạt động (TrangThai+còn hạn), Ngôn ngữ (3).",
        ],
        "operation": [
            "Index.OnGetAsync() nhận tham số mode/date/week/month/year/from/to rồi ResolveRange() để xác định SinceUtc, UntilUtc, Granularity và RangeLabel.",
            "LoadPoiSectionAsync() tải danh sách POI và tính TongPOI (tất cả POI), POIDangHoatDong (TrangThai=true và còn hạn duy trì), SoQuanQuaHan với cơ chế retry khi gặp lỗi disposed wait handle.",
            "LoadAnalyticsWithRetryAsync() gọi LoadActivityAsync(), LoadGeoAsync(), LoadTopPoiAsync(), LoadRevenueAsync() và LoadSummaryAsync() trên shared DbConnection.",
            "Sau khi có dữ liệu, hệ thống chạy FillMissingBuckets() rồi serialize ActivityJson, GeoJson và RevenueJson để giao diện JS dùng trực tiếp.",
            "Dashboard hiện tại không polling tự động; dữ liệu cập nhật theo mỗi lần admin đổi mốc thời gian hoặc tải lại trang.",
        ],
    },
]


SUB_FLOWS = [
    {
        "id": "F03b",
        "title": "Hàng đợi audio thuyết minh per-device",
        "parent": "F03",
        "activity_image": "03b-queue-playback-activity.png",
        "sequence_image": "03b-queue-playback-sequence.png",
        "function_text": (
            "Sub-flow F03b tách riêng cơ chế hàng đợi audio trong 1 thiết bị. Đây là queue per-device, "
            "không phải server-side queue. Khi nhiều khách cùng vào 1 quán, mỗi máy có queue riêng."
        ),
        "capabilities": [
            "Lớp 1 — _playingPoi: nếu POI đang phát chính là POI đang yêu cầu → skip.",
            "Lớp 2 — _queuedSpeakPoiIds (HashSet): nếu POI đã có trong queue → skip enqueue.",
            "Lớp 3 — _speakLock (SemaphoreSlim): serial hóa TTS, chỉ 1 SpeakAsync chạy 1 lúc.",
            "FIFO drain: khi đang phát POI A và khách vào POI B, B được enqueue và phát sau A xong.",
        ],
    },
]


# ──────────────────────────────────────────────────────────────────────────────
# Helpers (giữ + bổ sung)
# ──────────────────────────────────────────────────────────────────────────────

def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), "Times New Roman")
    existing = r_pr.find(qn("w:rFonts"))
    if existing is not None:
        r_pr.remove(existing)
    r_pr.insert(0, r_fonts)


def add_paragraph(doc, text, *, size=12, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
                  color=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13.5, 3: 12.5, 4: 11.8}
    colors = {1: (0, 70, 127), 2: (31, 73, 125), 3: (63, 63, 63), 4: (96, 96, 96)}
    p = add_paragraph(
        doc, text,
        size=sizes.get(level, 12),
        bold=True,
        color=colors.get(level),
        space_before=14 if level == 1 else (10 if level == 2 else 8),
        space_after=6 if level == 1 else 4,
    )
    if level == 1:
        # underline-style accent for level-1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.8 + (level * 0.5))
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=11.5)
    return paragraph


def add_table(doc, headers, rows, widths_cm=None, header_fill="1F497D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), header_fill)
        tc_pr.append(shading)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(header)
        set_run_font(run, size=11, bold=True, color=(255, 255, 255))

    for row_index, row in enumerate(rows, start=1):
        fill = "EAF1FB" if row_index % 2 == 1 else "FFFFFF"
        for col_index, value in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)
            tc_pr.append(shading)
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.space_before = Pt(2)
            cell_para.paragraph_format.space_after = Pt(2)
            run = cell_para.add_run(str(value))
            set_run_font(run, size=10.5)

    if widths_cm:
        for row in table.rows:
            for cell, width in zip(row.cells, widths_cm):
                cell.width = Cm(width)
    return table


def add_callout(doc, text, *, color=(31, 73, 125), bg="EAF1FB"):
    """Box màu thay cho blockquote — dùng cho ghi chú quan trọng."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.rows[0].cells[0]
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), bg)
    tc_pr.append(shading)
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run_font(run, size=11, italic=True, color=color)
    # spacing after the table
    add_paragraph(doc, "", size=2, space_after=2)
    return table


def add_image(doc, image_name, caption, width_cm):
    path = DIAGRAM_PNG_DIR / image_name
    if not path.exists():
        add_paragraph(doc, f"[Thiếu ảnh: {image_name}]", size=11, italic=True, color=(192, 0, 0))
        return
    caption_paragraph = add_paragraph(
        doc, caption,
        size=10.5, bold=True, italic=True, color=(96, 96, 96),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2,
    )
    caption_paragraph.paragraph_format.keep_with_next = True
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_after = Pt(10)
    picture_paragraph.add_run().add_picture(str(path), width=Cm(width_cm))


def add_page_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_left = para.add_run(f"PRD VinhKhanhTour v{PRODUCT_VERSION}  •  Tài liệu v{DOC_VERSION}  •  ")
    set_run_font(run_left, size=9, color=(120, 120, 120))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = para.add_run()
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_end)
    set_run_font(page_run, size=9, bold=True, color=(31, 73, 125))


# ──────────────────────────────────────────────────────────────────────────────
# Build document
# ──────────────────────────────────────────────────────────────────────────────

def build_cover(doc):
    add_paragraph(doc, "PRODUCT REQUIREMENTS DOCUMENT", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=80, space_after=4, color=(0, 70, 127))
    add_paragraph(doc, "VinhKhanhTour", size=28, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8, color=(31, 73, 125))
    add_paragraph(doc, "Hệ thống du lịch ẩm thực thông minh — phố Vĩnh Khánh, Quận 4, TP.HCM",
                  size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=0, space_after=40, color=(89, 89, 89))

    add_table(doc,
              ["Mục", "Giá trị"],
              [
                  ["Sản phẩm", f"VinhKhanhTour v{PRODUCT_VERSION}"],
                  ["Tài liệu", f"PRD v{DOC_VERSION}"],
                  ["Ngày phát hành", DOC_DATE],
                  ["Phạm vi", "MAUI App + REST API + CMS Razor Pages"],
                  ["Nguồn sơ đồ", "docs/diagrams/*.puml (PlantUML)"],
                  ["Trạng thái", "Final — Sẵn sàng phát hành"],
              ],
              widths_cm=[5.0, 10.5])

    add_paragraph(doc, "", space_after=20)
    add_paragraph(
        doc,
        "Tài liệu này mô tả toàn bộ phạm vi, mục tiêu, đối tượng người dùng, yêu cầu chức năng và phi chức năng "
        "của hệ thống VinhKhanhTour. Đây là nguồn tham chiếu chính cho đội phát triển, kiểm thử và quản lý sản phẩm.",
        size=11.5, italic=True, color=(89, 89, 89),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8,
    )
    doc.add_page_break()


def build_version_history(doc):
    add_heading(doc, "Lịch sử phiên bản tài liệu", level=1)
    add_table(doc,
              ["Version", "Ngày", "Mô tả thay đổi"],
              VERSION_HISTORY,
              widths_cm=[2.5, 3.2, 9.8])
    doc.add_page_break()


def build_toc(doc):
    add_heading(doc, "Mục lục", level=1)
    sections = [
        ("PHẦN A — SẢN PHẨM", None),
        ("1. Tóm tắt điều hành", "Executive summary, vấn đề và giải pháp."),
        ("2. Bối cảnh & Vấn đề", "Bối cảnh kinh doanh, các vấn đề cần giải quyết."),
        ("3. Mục tiêu sản phẩm & KPI", "Mục tiêu cụ thể và chỉ số đo lường thành công."),
        ("4. Đối tượng người dùng (Personas)", "Khách quốc tế, khách trong nước, admin, chủ quán."),
        ("5. Phạm vi sản phẩm", "Trong & ngoài phạm vi phiên bản 1.5."),
        ("6. Mô hình kinh doanh & Bảng giá", "Gói app khách + phí duy trì POI."),
        ("PHẦN B — TRẢI NGHIỆM", None),
        ("7. User journey", "Hành trình khách + admin từ đầu đến cuối."),
        ("8. Use case tổng thể", "Sơ đồ use case của hệ thống."),
        ("PHẦN C — KỸ THUẬT", None),
        ("9. Kiến trúc hệ thống", "MAUI ↔ API ↔ CMS ↔ DB."),
        ("10. Mô hình dữ liệu (ERD)", "11 bảng và quan hệ."),
        ("11. State machine", "YeuCauThanhToan: cho_duyet → da_duyet/tu_choi."),
        ("12. Hằng số thời gian quan trọng", "GPS poll, heartbeat, dwell, cooldown..."),
        ("PHẦN D — CHỨC NĂNG", None),
        ("13. Phân tích chi tiết F01 đến F10", "Mô tả + activity + sequence cho mỗi feature."),
        ("14. Sub-flow F03b — Hàng đợi audio", "Per-device queue với 3 lớp chống trùng."),
        ("PHẦN E — PHI CHỨC NĂNG", None),
        ("15. Yêu cầu phi chức năng (NFR)", "Hiệu năng, độ tin cậy, bảo mật, i18n."),
        ("PHẦN F — API REFERENCE", None),
        ("16. Tổng hợp endpoint REST API", "POI, Subscription, Heartbeat, Payment, Khác."),
        ("PHẦN G — VẬN HÀNH", None),
        ("17. Rủi ro & Giảm thiểu", "7 rủi ro chính + đề xuất xử lý."),
        ("18. Tiêu chí chấp nhận (DoD)", "Tiêu chí pass khi release."),
        ("19. Kiểm thử tự động & Testability", "API, CMS, Appium smoke tests và script tổng hợp."),
        ("20. Giả định & Ràng buộc", "Tiền đề và giới hạn kỹ thuật."),
        ("21. Lịch trình & Cột mốc", "Roadmap từ v1.0 đến v2.0."),
        ("PHẦN H — PHỤ LỤC", None),
        ("22. Thuật ngữ (Glossary)", "Định nghĩa các khái niệm chuyên ngành."),
        ("23. Tham chiếu tài liệu", "Sơ đồ PlantUML, file mã nguồn liên quan."),
    ]

    for title, desc in sections:
        if desc is None:
            add_paragraph(doc, title, size=12, bold=True, color=(31, 73, 125),
                          space_before=10, space_after=4)
        else:
            row = doc.add_paragraph()
            row.paragraph_format.left_indent = Cm(0.6)
            row.paragraph_format.space_after = Pt(3)
            run_title = row.add_run(title)
            set_run_font(run_title, size=11, bold=True)
            run_desc = row.add_run(f"  —  {desc}")
            set_run_font(run_desc, size=10.5, italic=True, color=(96, 96, 96))
    doc.add_page_break()


def build_part_a(doc):
    add_paragraph(doc, "PHẦN A — SẢN PHẨM", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    # 1. Executive summary
    add_heading(doc, "1. Tóm tắt điều hành", level=1)
    add_paragraph(doc, EXECUTIVE_SUMMARY, size=11.5)
    add_callout(
        doc,
        "Điểm khác biệt: hệ thống tự động phát thuyết minh đa ngôn ngữ theo vị trí GPS, cho phép khách "
        "dạo phố rảnh tay; đồng thời mô hình hai chiều (gói khách + phí duy trì POI) tạo nguồn thu bền vững."
    )

    # 2. Background
    add_heading(doc, "2. Bối cảnh & Vấn đề", level=1)
    add_paragraph(doc,
                  "Phố ẩm thực Vĩnh Khánh (Quận 4, TP.HCM) là tuyến phố ẩm thực nổi tiếng với hàng chục quán "
                  "san sát, thu hút lượng khách lớn — đặc biệt khách quốc tế. Tuy nhiên, các vấn đề sau làm "
                  "giảm trải nghiệm và doanh thu:", size=11.5)
    for item in PROBLEM_STATEMENT:
        add_bullet(doc, item)

    # 3. Goals
    add_heading(doc, "3. Mục tiêu sản phẩm & KPI", level=1)
    add_paragraph(doc,
                  "Năm mục tiêu cốt lõi và chỉ số đo lường thành công cho phiên bản 1.5:", size=11.5)
    add_table(doc,
              ["Mục tiêu", "Mô tả", "KPI / Ngưỡng chấp nhận"],
              GOALS_KPI,
              widths_cm=[4.5, 6.5, 5.0])

    # 4. Personas
    add_heading(doc, "4. Đối tượng người dùng (Personas)", level=1)
    for persona in PERSONAS:
        add_heading(doc, f"{persona['name']}  ({persona['role']})", level=3)
        add_table(doc, ["Mục", "Nội dung"],
                  [
                      ["Đặc điểm", persona["demographics"]],
                      ["Thiết bị / Công nghệ", persona["tech"]],
                  ], widths_cm=[4.0, 11.0])
        add_paragraph(doc, "Mục tiêu chính:", size=11, bold=True, space_before=4, space_after=2)
        for g in persona["goals"]:
            add_bullet(doc, g)
        add_paragraph(doc, "Vấn đề / Pain points:", size=11, bold=True, space_before=4, space_after=2)
        for p in persona["pain"]:
            add_bullet(doc, p)

    # 5. Scope
    add_heading(doc, "5. Phạm vi sản phẩm", level=1)
    add_heading(doc, "Trong phạm vi (In scope)", level=3)
    for item in IN_SCOPE:
        add_bullet(doc, item)
    add_heading(doc, "Ngoài phạm vi (Out of scope)", level=3)
    for item in OUT_OF_SCOPE:
        add_bullet(doc, item)

    # 6. Pricing
    add_heading(doc, "6. Mô hình kinh doanh & Bảng giá", level=1)
    add_heading(doc, "6.1 Gói đăng ký App (khách tham quan)", level=3)
    add_table(doc, ["Mã gói", "Tên", "Giá", "Thời hạn", "Ghi chú"], PRICING_APP,
              widths_cm=[1.8, 2.2, 2.5, 2.5, 6.0])
    add_callout(
        doc,
        "Gói nối tiếp: khi gia hạn lúc gói cũ chưa hết hạn, NgayBatDau = NgayHetHan cũ (cộng dồn, không mất ngày)."
    )

    add_heading(doc, "6.2 Phí dịch vụ POI (chủ quán)", level=3)
    add_table(doc, ["Loại phí", "Mức phí", "Chu kỳ", "Ghi chú"], PRICING_POI,
              widths_cm=[4.0, 2.5, 2.5, 6.0])
    doc.add_page_break()


def build_part_b(doc):
    add_paragraph(doc, "PHẦN B — TRẢI NGHIỆM", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    # 7. User journey
    add_heading(doc, "7. User journey", level=1)
    add_heading(doc, "7.1 Hành trình khách du lịch", level=3)
    for step in USER_JOURNEY_GUEST:
        add_bullet(doc, step)
    add_heading(doc, "7.2 Hành trình quản trị viên CMS", level=3)
    for step in USER_JOURNEY_ADMIN:
        add_bullet(doc, step)

    # 8. Use case overall
    add_heading(doc, "8. Use case tổng thể", level=1)
    add_paragraph(doc, USE_CASE_SUMMARY["description"], size=11.5)
    add_paragraph(doc, USE_CASE_SUMMARY["coverage_note"], size=11.5)
    add_image(doc, USE_CASE_SUMMARY["image"], "Hình 8.1 — Use case tổng thể của hệ thống", 15.8)
    doc.add_page_break()


def build_part_c(doc):
    add_paragraph(doc, "PHẦN C — KỸ THUẬT", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    # 9. Architecture
    add_heading(doc, "9. Kiến trúc hệ thống", level=1)
    add_paragraph(doc,
                  "Hệ thống gồm 3 thành phần chính giao tiếp với nhau qua HTTP và một database tập trung "
                  "trên Supabase. Mọi dữ liệu nghiệp vụ đều đi qua REST API.", size=11.5)
    add_table(doc,
              ["Thành phần", "Công nghệ", "Vai trò"],
              [
                  ["Mobile App", ".NET MAUI 10 (Android)", "UI khách: subscription, POI, geofence audio, thanh toán."],
                  ["REST API", "ASP.NET Core 10 + EF Core", "Backend trung tâm; xử lý nghiệp vụ và lưu DB."],
                  ["CMS Web", "ASP.NET Core Razor Pages", "Quản trị nội dung, thanh toán, theo dõi khách."],
                  ["Database", "Supabase PostgreSQL (managed)", "Lưu toàn bộ dữ liệu nghiệp vụ."],
                  ["Hosting API", "Render.com (Singapore region)", "Public URL: phoamthuc.onrender.com"],
              ], widths_cm=[3.5, 4.5, 7.0])
    add_heading(doc, "Ghi chú kiến trúc", level=3)
    for item in ARCHITECTURE_NOTES:
        add_bullet(doc, item)

    # 10. ERD
    add_heading(doc, "10. Mô hình dữ liệu (ERD)", level=1)
    add_paragraph(doc,
                  "Database gồm 11 bảng, tổ chức theo 3 vùng: nội dung POI, phí duy trì, và khách hàng ẩn danh. "
                  "Các bảng vùng khách (dangkyapp / yeucauthanhtoan / vitrikhach / lichsuphat) liên kết với nhau "
                  "qua MaThietBi (logical key, không có FK constraint).", size=11.5)
    add_image(doc, "12-erd.png", "Hình 10.1 — ERD tổng thể", 16.0)
    add_table(doc, ["Bảng", "Mô tả", "Cột quan trọng"], DATA_TABLES,
              widths_cm=[3.0, 5.0, 7.0])

    # 11. State machine
    add_heading(doc, "11. State machine — YeuCauThanhToan", level=1)
    add_paragraph(doc,
                  "Yêu cầu thanh toán đi qua 3 trạng thái: cho_duyet (mặc định khi tạo) → "
                  "da_duyet (admin Approve) hoặc tu_choi (admin Reject). Cả 2 trạng thái cuối là terminal.",
                  size=11.5)
    add_image(doc, "11-yeucauthanhtoan-state.png", "Hình 11.1 — State machine YeuCauThanhToan", 14.5)

    # 12. Timing
    add_heading(doc, "12. Hằng số thời gian quan trọng", level=1)
    add_paragraph(doc,
                  "Bảng dưới tổng hợp các chu kỳ và ngưỡng thời gian được hard-code trong code, "
                  "đóng vai trò xương sống cho hành vi realtime của app và CMS.", size=11.5)
    add_table(doc, ["Tham số", "Giá trị", "Nguồn / Mô tả"], TIMING_CONSTANTS,
              widths_cm=[4.5, 3.0, 7.5])
    doc.add_page_break()


def build_part_d(doc):
    add_paragraph(doc, "PHẦN D — CHỨC NĂNG", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    add_heading(doc, "13. Phân tích chi tiết F01 đến F10", level=1)
    add_paragraph(doc,
                  "Mỗi chức năng được mô tả với: (a) actor chính, (b) tính năng cốt lõi, "
                  "(c) cách hoạt động theo code thật, (d) sơ đồ activity và sequence được render từ PlantUML.",
                  size=11.5)
    add_callout(
        doc,
        "Cách đọc: Activity diagram cho biết LUỒNG xử lý nội bộ; Sequence diagram cho biết THỨ TỰ tương tác "
        "giữa UI ↔ controller ↔ database. Đọc cùng nhau để có bức tranh đầy đủ."
    )

    for index, feature in enumerate(FEATURES, start=1):
        doc.add_page_break()
        add_heading(doc, f"13.{index}  {feature['id']} — {feature['title']}", level=2)
        add_table(doc, ["Mục", "Nội dung"],
                  [
                      ["Actor chính", feature["actor"]],
                      ["Trạng thái", feature["status"]],
                      ["Sơ đồ activity", feature["activity_image"]],
                      ["Sơ đồ sequence", feature["sequence_image"]],
                  ], widths_cm=[3.5, 11.5])
        add_heading(doc, "Mô tả chức năng", level=3)
        add_paragraph(doc, feature["function_text"], size=11.5)
        add_heading(doc, "Tính năng cốt lõi", level=3)
        for item in feature["capabilities"]:
            add_bullet(doc, item)
        if feature.get("implementation"):
            add_heading(doc, "Phân lớp triển khai", level=3)
            add_table(
                doc,
                ["Lớp", "Thành phần / method names"],
                feature["implementation"],
                widths_cm=[3.4, 11.6])
        add_heading(doc, "Cách hoạt động theo code", level=3)
        for item in feature["operation"]:
            add_bullet(doc, item)
        add_image(doc, feature["activity_image"],
                  f"Hình 13.{index}a — {feature['id']} Activity", 14.0)
        add_image(doc, feature["sequence_image"],
                  f"Hình 13.{index}b — {feature['id']} Sequence", 15.6)

    # Sub-flows
    doc.add_page_break()
    add_heading(doc, "14. Sub-flow F03b — Hàng đợi audio per-device", level=1)
    sub = SUB_FLOWS[0]
    add_paragraph(doc, sub["function_text"], size=11.5)
    add_heading(doc, "Cơ chế chống trùng (3 lớp)", level=3)
    for item in sub["capabilities"]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Lưu ý: đây là queue PER-DEVICE. Khi nhiều khách cùng vào một quán, mỗi máy có queue riêng — "
        "không có server-side queue. Việc chống lặp giữa các thiết bị dựa vào visit dedup 10 phút ở DB."
    )
    add_image(doc, sub["activity_image"], "Hình 14.1 — F03b Activity (queue + drain + speak)", 14.0)
    add_image(doc, sub["sequence_image"], "Hình 14.2 — F03b Sequence (4 tình huống thực tế)", 15.6)
    doc.add_page_break()


def build_part_e(doc):
    add_paragraph(doc, "PHẦN E — PHI CHỨC NĂNG", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    add_heading(doc, "15. Yêu cầu phi chức năng (NFR)", level=1)

    add_heading(doc, "15.1 Hiệu năng", level=3)
    add_table(doc, ["Chỉ số", "Mục tiêu", "Ghi chú"], NFR_PERFORMANCE,
              widths_cm=[4.5, 4.0, 6.5])

    add_heading(doc, "15.2 Độ tin cậy", level=3)
    for item in NFR_RELIABILITY:
        add_bullet(doc, item)

    add_heading(doc, "15.3 Bảo mật & Quyền riêng tư", level=3)
    for item in NFR_SECURITY:
        add_bullet(doc, item)

    add_heading(doc, "15.4 Khả năng mở rộng & Quốc tế hoá", level=3)
    for item in I18N_NOTES:
        add_bullet(doc, item)
    doc.add_page_break()


def build_part_f(doc):
    add_paragraph(doc, "PHẦN F — API REFERENCE", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    add_heading(doc, "16. Tổng hợp endpoint REST API", level=1)
    add_paragraph(doc,
                  f"Base URL: https://phoamthuc.onrender.com  |  Phiên bản API: v{PRODUCT_VERSION}",
                  size=11, italic=True, color=(89, 89, 89))

    for group, endpoints in API_ENDPOINTS.items():
        add_heading(doc, f"16.{list(API_ENDPOINTS).index(group) + 1}  {group}", level=3)
        add_table(doc, ["Method", "Endpoint", "Mô tả"], endpoints,
                  widths_cm=[1.8, 6.2, 7.0])
    doc.add_page_break()


def build_part_g(doc):
    add_paragraph(doc, "PHẦN G — VẬN HÀNH", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    add_heading(doc, "17. Rủi ro & Giảm thiểu", level=1)
    add_table(doc, ["Mã", "Mức", "Mô tả rủi ro", "Đề xuất giảm thiểu"], RISKS,
              widths_cm=[1.6, 1.8, 6.0, 5.6])

    add_heading(doc, "18. Tiêu chí chấp nhận (Definition of Done)", level=1)
    add_paragraph(doc,
                  "Phiên bản phát hành chỉ được coi là sẵn sàng khi tất cả tiêu chí dưới đây đều đạt:",
                  size=11.5)
    for item in ACCEPTANCE_CRITERIA:
        add_bullet(doc, item)

    add_heading(doc, "19. Kiểm thử tự động & Testability", level=1)
    add_paragraph(doc,
                  "Baseline kiểm thử hiện tại tập trung vào smoke/integration tests để xác nhận API, CMS và môi trường mobile automation "
                  "sẵn sàng trước khi demo hoặc phát hành nội bộ.",
                  size=11.5)
    add_table(doc,
              ["Hạng mục", "Công nghệ", "Phạm vi kiểm tra"],
              TEST_AUTOMATION,
              widths_cm=[4.0, 4.5, 6.5])
    add_heading(doc, "Ghi chú vận hành test", level=3)
    for item in TESTABILITY_NOTES:
        add_bullet(doc, item)

    add_heading(doc, "20. Giả định & Ràng buộc", level=1)
    add_heading(doc, "20.1 Giả định", level=3)
    for item in ASSUMPTIONS:
        add_bullet(doc, item)
    add_heading(doc, "20.2 Ràng buộc", level=3)
    for item in CONSTRAINTS:
        add_bullet(doc, item)

    add_heading(doc, "21. Lịch trình & Cột mốc", level=1)
    add_table(doc, ["Phiên bản", "Mốc thời gian", "Nội dung"], ROADMAP,
              widths_cm=[3.5, 3.0, 8.5])
    doc.add_page_break()


def build_part_h(doc):
    add_paragraph(doc, "PHẦN H — PHỤ LỤC", size=14, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 73, 125),
                  space_before=10, space_after=14)

    add_heading(doc, "22. Thuật ngữ (Glossary)", level=1)
    add_table(doc, ["Thuật ngữ", "Định nghĩa"], GLOSSARY,
              widths_cm=[4.0, 11.0])

    add_heading(doc, "23. Tham chiếu tài liệu", level=1)
    add_heading(doc, "23.1 Sơ đồ PlantUML", level=3)
    diagram_refs = [
        ["00", "Use case tổng thể", "00-overall-usecase.puml"],
        ["F01", "Subscription gate", "01-subscription-gate-{activity,sequence}.puml"],
        ["F02", "Khám phá POI", "02-poi-explore-{activity,sequence}.puml"],
        ["F03", "Geofence + audio", "03-geofence-audio-{activity,sequence}.puml"],
        ["F03b", "Hàng đợi audio per-device", "03b-queue-playback-{activity,sequence}.puml"],
        ["F04", "Chi tiết POI", "04-poi-detail-{activity,sequence}.puml"],
        ["F05", "Thanh toán gói", "05-paid-plan-{activity,sequence}.puml"],
        ["F06", "Quản lý POI CMS", "06-cms-poi-management-{activity,sequence}.puml"],
        ["F07", "Phí duy trì POI", "07-maintenance-payment-{activity,sequence}.puml"],
        ["F08", "Duyệt thanh toán", "08-app-payment-approval-{activity,sequence}.puml"],
        ["F09", "BanDo live tracking", "09-live-map-{activity,sequence}.puml"],
        ["F10", "Dashboard CMS", "10-dashboard-{activity,sequence}.puml"],
        ["F11", "State YeuCauThanhToan", "11-yeucauthanhtoan-state.puml"],
        ["F12", "ERD tổng thể", "12-erd.puml"],
    ]
    add_table(doc, ["ID", "Tên sơ đồ", "File nguồn"], diagram_refs,
              widths_cm=[1.5, 5.5, 8.0])

    add_heading(doc, "23.2 Mã nguồn liên quan", level=3)
    code_refs = [
        ["VinhKhanhTourDemo/MainPage.xaml.cs", "Vòng lặp GPS, geofence, queue audio."],
        ["VinhKhanhTourDemo/SubscriptionPage.xaml.cs", "Gate gói + recovery code/QR."],
        ["VinhKhanhTourDemo/PaymentPage.xaml.cs", "QR VietQR + nội dung CK."],
        ["VinhKhanhTour.API/Controllers/HeartbeatController.cs", "SendHeartbeat + verify POI."],
        ["VinhKhanhTour.API/Controllers/SubscriptionController.cs", "CreateRequest + CRUD gói."],
        ["VinhKhanhTour.CMS/Pages/Index.cshtml(.cs)", "Dashboard 3 thẻ + analytics."],
        ["VinhKhanhTour.CMS/Pages/DuyetThanhToan/", "Approve/Reject + snapshot polling."],
        ["VinhKhanhTour.CMS/Pages/BanDo/", "Live map + customer tracking."],
        ["tests/VinhKhanhTour.API.Tests/", "xUnit smoke/integration test cho API health endpoint."],
        ["tests/VinhKhanhTour.CMS.E2ETests/", "Playwright smoke/E2E test cho CMS health + Privacy page."],
        ["tests/VinhKhanhTour.MAUI.AppiumTests/", "NUnit/Appium opt-in smoke test cho mobile automation server."],
        ["scripts/run-all-tests.ps1", "Script restore/build/install Playwright/run automated smoke tests."],
    ]
    add_table(doc, ["File / Thư mục", "Vai trò"], code_refs,
              widths_cm=[7.5, 7.5])

    add_paragraph(doc, "", space_after=8)
    add_paragraph(doc,
                  f"— Hết —  |  PRD VinhKhanhTour v{PRODUCT_VERSION} — Tài liệu v{DOC_VERSION} — {DOC_DATE}",
                  size=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120, 120, 120),
                  space_before=24)


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.0)

    add_page_footer(doc)

    build_cover(doc)
    build_version_history(doc)
    build_toc(doc)
    build_part_a(doc)
    build_part_b(doc)
    build_part_c(doc)
    build_part_d(doc)
    build_part_e(doc)
    build_part_f(doc)
    build_part_g(doc)
    build_part_h(doc)

    return doc


def main():
    document = build_document()
    document.save(str(OUT_DOCX))
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()
